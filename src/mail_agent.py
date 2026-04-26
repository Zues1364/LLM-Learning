import base64
import hashlib
import hmac
import json
import logging
import os
import re
import sqlite3
import secrets
import time
from datetime import datetime, timedelta, timezone
from io import BytesIO
from pathlib import Path
from threading import Lock
from typing import Any, Dict, List, Optional, Tuple
from urllib.parse import urlencode, urlparse
from uuid import uuid4

import requests

from utils import normalize_for_match

logger = logging.getLogger(__name__)

BASE_DIR = Path(__file__).resolve().parent.parent
MAIL_DIR = BASE_DIR / "data" / "mail"
MAIL_SESSION_DIR = MAIL_DIR / "sessions"
MAIL_OAUTH_STATE_FILE = MAIL_DIR / "oauth_states.json"
MAIL_DB_PATH = BASE_DIR / "data" / "memory.db"
RESOURCE_SESSION_DIR = BASE_DIR / "data" / "resources" / "sessions"
RESOURCE_USER_DIR = BASE_DIR / "data" / "resources" / "users"

SUPPORTED_EXTENSIONS = {".pdf", ".html", ".htm", ".docx", ".xlsx", ".xls"}
DEFAULT_KEYWORDS = [
    "mo lop",
    "mở lớp",
    "dang ky hoc phan",
    "đăng ký học phần",
    "hoc ky",
    "học kỳ",
    "thoi khoa bieu",
    "thời khóa biểu",
    "tkb",
    "lop hoc",
]
DEFAULT_SCHEDULE_TOKENS = [
    "thoi khoa bieu",
    "tkb",
    "phu luc",
    "cong van",
    "lich hoc",
    "lop hoc",
    "mo lop",
]
DEFAULT_REGISTRATION_TOKENS = [
    "dang ky hoc phan",
    "dang ky hoc",
    "mo cong",
    "dkmh",
    "tin chi",
    "hoc ky",
    "phien thu nhat",
    "phien thu hai",
]
DEFAULT_NEGATIVE_TOKENS = [
    "job fair",
    "hoi cho viec lam",
    "tuyen dung",
    "su kien",
    "workshop",
    "cuoc thi",
    "hoc bong",
]


def _utc_now() -> datetime:
    return datetime.now(timezone.utc)


def _utc_now_iso() -> str:
    return _utc_now().isoformat()


def _safe_session_id(session_id: str) -> str:
    text = (session_id or "default").strip()
    text = re.sub(r"[^A-Za-z0-9._-]", "_", text)
    return text or "default"


def _safe_owner_id(value: str) -> str:
    text = (value or "").strip()
    text = re.sub(r"[^A-Za-z0-9._-]", "_", text)
    return text or "default"


def _read_json(path: Path, default: Any) -> Any:
    try:
        if not path.exists():
            return default
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return default


def _write_json(path: Path, payload: Any):
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def _decode_b64url(data: str) -> str:
    if not data:
        return ""
    try:
        raw = data.encode("utf-8")
        padding = b"=" * ((4 - len(raw) % 4) % 4)
        decoded = base64.urlsafe_b64decode(raw + padding)
        try:
            return decoded.decode("utf-8")
        except UnicodeDecodeError:
            return decoded.decode("latin-1", errors="ignore")
    except Exception:
        return ""


def _decode_b64url_bytes(data: str) -> bytes:
    if not data:
        return b""
    raw = data.encode("utf-8")
    padding = b"=" * ((4 - len(raw) % 4) % 4)
    return base64.urlsafe_b64decode(raw + padding)


def _extract_email(value: str) -> str:
    text = value or ""
    m = re.search(r"<([^>]+)>", text)
    if m:
        return m.group(1).strip().lower()
    if "@" in text:
        return text.strip().lower()
    return ""


def _extract_domain(email: str) -> str:
    if "@" not in email:
        return ""
    return email.split("@", 1)[1].strip().lower()


def _iter_payload_parts(node: Dict[str, Any]) -> List[Dict[str, Any]]:
    parts: List[Dict[str, Any]] = []
    stack = [node or {}]
    while stack:
        current = stack.pop()
        if not isinstance(current, dict):
            continue
        parts.append(current)
        for child in current.get("parts") or []:
            stack.append(child)
    return parts


def _normalize_keywords(raw: str) -> List[str]:
    if not raw:
        return list(DEFAULT_KEYWORDS)
    values = [item.strip() for item in raw.split(",") if item.strip()]
    return values or list(DEFAULT_KEYWORDS)


def _normalize_csv_tokens(raw: str, defaults: List[str]) -> List[str]:
    if not raw:
        return list(defaults)
    values = [item.strip() for item in raw.split(",") if item.strip()]
    return values or list(defaults)


def _clip01(value: float) -> float:
    return max(0.0, min(1.0, value))


def _safe_float(raw: str, default: float) -> float:
    try:
        return float(raw)
    except Exception:
        return default


class MailAgentService:
    def __init__(self):
        self.gmail_client_id = os.getenv("GOOGLE_OAUTH_CLIENT_ID", "").strip()
        self.gmail_client_secret = os.getenv("GOOGLE_OAUTH_CLIENT_SECRET", "").strip()
        self.gmail_redirect_default = os.getenv("MAIL_OAUTH_REDIRECT_URI", "").strip()
        self.app_auth_redirect_default = (
            os.getenv("APP_OAUTH_REDIRECT_URI", "").strip()
            or "http://127.0.0.1:9000/api/auth/google/callback"
        )
        self.app_auth_scope = (
            os.getenv("APP_GOOGLE_AUTH_SCOPE", "openid email profile").strip() or "openid email profile"
        )
        self.gmail_scope = os.getenv(
            "MAIL_GMAIL_SCOPE",
            "https://www.googleapis.com/auth/gmail.readonly",
        ).strip()
        self.poll_minutes = max(1, int(os.getenv("MAIL_POLL_MINUTES", "5") or "5"))
        self.relevance_keywords = _normalize_keywords(os.getenv("MAIL_RELEVANCE_KEYWORDS", ""))
        self.trusted_domains = [
            d.strip().lower()
            for d in (os.getenv("MAIL_TRUSTED_DOMAINS", "uet.edu.vn,vnu.edu.vn").split(","))
            if d.strip()
        ]
        self.query = os.getenv("MAIL_GMAIL_QUERY", "newer_than:7d").strip() or "newer_than:7d"
        self.retention_days = max(1, int(os.getenv("MAIL_RETENTION_DAYS", "7") or "7"))
        self.gemini_api_key = os.getenv("GEMINI_API_KEY", "").strip()
        self.intent_mode = (os.getenv("MAIL_INTENT_CLASSIFIER_MODE", "hybrid") or "hybrid").strip().lower()
        if self.intent_mode not in {"rule_only", "llm_only", "hybrid"}:
            self.intent_mode = "hybrid"
        self.intent_llm_model = (
            os.getenv("MAIL_INTENT_LLM_MODEL", "gemini-2.5-flash") or "gemini-2.5-flash"
        ).strip()
        self.intent_llm_threshold = _clip01(
            _safe_float(os.getenv("MAIL_INTENT_LLM_THRESHOLD", "0.70") or "0.70", 0.70)
        )
        self.oauth_state_ttl_seconds = max(
            60, int(os.getenv("MAIL_OAUTH_STATE_TTL_SECONDS", "600") or "600")
        )
        self.app_session_cookie_name = (
            os.getenv("APP_SESSION_COOKIE_NAME", "rag_cosmic_session").strip() or "rag_cosmic_session"
        )
        self.app_session_ttl_days = max(1, int(os.getenv("APP_SESSION_TTL_DAYS", "7") or "7"))
        self.app_session_secret = (
            os.getenv("APP_SESSION_SECRET", "").strip()
            or self.gmail_client_secret
            or "rag-cosmic-dev-session-secret"
        )
        self.intent_schedule_tokens = _normalize_csv_tokens(
            os.getenv("MAIL_SCHEDULE_TOKENS", ""),
            DEFAULT_SCHEDULE_TOKENS,
        )
        self.intent_registration_tokens = _normalize_csv_tokens(
            os.getenv("MAIL_REGISTRATION_TOKENS", ""),
            DEFAULT_REGISTRATION_TOKENS,
        )
        self.intent_negative_tokens = _normalize_csv_tokens(
            os.getenv("MAIL_NEGATIVE_TOKENS", ""),
            DEFAULT_NEGATIVE_TOKENS,
        )

        MAIL_DIR.mkdir(parents=True, exist_ok=True)
        MAIL_SESSION_DIR.mkdir(parents=True, exist_ok=True)
        RESOURCE_SESSION_DIR.mkdir(parents=True, exist_ok=True)
        RESOURCE_USER_DIR.mkdir(parents=True, exist_ok=True)
        self._owner_locks: Dict[str, Lock] = {}
        self._owner_locks_guard = Lock()
        self._ensure_db_schema()

    def _db_conn(self) -> sqlite3.Connection:
        conn = sqlite3.connect(str(MAIL_DB_PATH), timeout=30, check_same_thread=False)
        conn.row_factory = sqlite3.Row
        return conn

    def _ensure_db_schema(self):
        with self._db_conn() as conn:
            conn.executescript(
                """
                CREATE TABLE IF NOT EXISTS users (
                    id TEXT PRIMARY KEY,
                    google_sub TEXT NOT NULL UNIQUE,
                    email TEXT NOT NULL UNIQUE,
                    name TEXT,
                    picture_url TEXT,
                    created_at TEXT NOT NULL,
                    updated_at TEXT NOT NULL
                );
                CREATE TABLE IF NOT EXISTS auth_sessions (
                    token_hash TEXT PRIMARY KEY,
                    user_id TEXT NOT NULL,
                    expires_at TEXT NOT NULL,
                    created_at TEXT NOT NULL,
                    last_seen_at TEXT NOT NULL
                );
                CREATE INDEX IF NOT EXISTS idx_auth_sessions_user ON auth_sessions(user_id);
                CREATE TABLE IF NOT EXISTS mail_connections (
                    user_id TEXT PRIMARY KEY,
                    gmail_email TEXT NOT NULL,
                    refresh_token_enc TEXT NOT NULL,
                    access_token TEXT,
                    access_expiry INTEGER,
                    scope TEXT NOT NULL,
                    connected_at TEXT NOT NULL,
                    updated_at TEXT NOT NULL
                );
                CREATE TABLE IF NOT EXISTS mail_migration_log (
                    id TEXT PRIMARY KEY,
                    user_id TEXT NOT NULL,
                    source_session_id TEXT NOT NULL,
                    status TEXT NOT NULL,
                    details_json TEXT,
                    created_at TEXT NOT NULL,
                    UNIQUE(user_id, source_session_id)
                );
                CREATE TABLE IF NOT EXISTS mail_user_state (
                    user_id TEXT PRIMARY KEY,
                    whitelist_json TEXT NOT NULL,
                    processed_message_ids_json TEXT NOT NULL,
                    candidates_json TEXT NOT NULL,
                    last_poll_at TEXT,
                    updated_at TEXT NOT NULL
                );
                """
            )
            conn.commit()

    def _hash_app_session_token(self, raw_token: str) -> str:
        return hashlib.sha256(
            f"{self.app_session_secret}:{raw_token}".encode("utf-8")
        ).hexdigest()

    def _cipher_keystream(self, nonce: bytes, size: int) -> bytes:
        out = bytearray()
        counter = 0
        secret = self.app_session_secret.encode("utf-8")
        while len(out) < size:
            out.extend(hashlib.sha256(secret + nonce + counter.to_bytes(4, "big")).digest())
            counter += 1
        return bytes(out[:size])

    def _encrypt_secret(self, raw: str) -> str:
        data = (raw or "").encode("utf-8")
        nonce = secrets.token_bytes(16)
        stream = self._cipher_keystream(nonce, len(data))
        payload = nonce + bytes(a ^ b for a, b in zip(data, stream))
        return base64.urlsafe_b64encode(payload).decode("ascii")

    def _decrypt_secret(self, payload: str) -> str:
        if not payload:
            return ""
        raw = base64.urlsafe_b64decode(payload.encode("ascii"))
        nonce, cipher = raw[:16], raw[16:]
        stream = self._cipher_keystream(nonce, len(cipher))
        plain = bytes(a ^ b for a, b in zip(cipher, stream))
        return plain.decode("utf-8")

    def _serialize_user_row(self, row: sqlite3.Row) -> Dict[str, Any]:
        return {
            "id": str(row["id"]),
            "email": str(row["email"]),
            "name": str(row["name"] or ""),
            "picture_url": str(row["picture_url"] or ""),
        }

    def _upsert_user(self, profile: Dict[str, Any]) -> Dict[str, Any]:
        google_sub = str(profile.get("sub") or "").strip()
        email = str(profile.get("email") or "").strip().lower()
        if not google_sub or not email:
            raise ValueError("Google profile missing sub/email.")
        now_iso = _utc_now_iso()
        with self._db_conn() as conn:
            existing = conn.execute(
                "SELECT * FROM users WHERE google_sub = ? OR email = ?",
                (google_sub, email),
            ).fetchone()
            if existing:
                user_id = str(existing["id"])
                conn.execute(
                    """
                    UPDATE users
                    SET google_sub = ?, email = ?, name = ?, picture_url = ?, updated_at = ?
                    WHERE id = ?
                    """,
                    (
                        google_sub,
                        email,
                        str(profile.get("name") or existing["name"] or ""),
                        str(profile.get("picture") or existing["picture_url"] or ""),
                        now_iso,
                        user_id,
                    ),
                )
            else:
                user_id = str(uuid4())
                conn.execute(
                    """
                    INSERT INTO users (id, google_sub, email, name, picture_url, created_at, updated_at)
                    VALUES (?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        user_id,
                        google_sub,
                        email,
                        str(profile.get("name") or ""),
                        str(profile.get("picture") or ""),
                        now_iso,
                        now_iso,
                    ),
                )
            conn.commit()
            row = conn.execute("SELECT * FROM users WHERE id = ?", (user_id,)).fetchone()
        if row is None:
            raise ValueError("Failed to upsert user.")
        return self._serialize_user_row(row)

    def _create_app_session(self, user_id: str) -> str:
        raw_token = secrets.token_urlsafe(32)
        token_hash = self._hash_app_session_token(raw_token)
        now_iso = _utc_now_iso()
        expires_at = (_utc_now() + timedelta(days=self.app_session_ttl_days)).isoformat()
        with self._db_conn() as conn:
            conn.execute(
                """
                INSERT OR REPLACE INTO auth_sessions (token_hash, user_id, expires_at, created_at, last_seen_at)
                VALUES (?, ?, ?, ?, ?)
                """,
                (token_hash, user_id, expires_at, now_iso, now_iso),
            )
            conn.commit()
        return raw_token

    def get_authenticated_user(self, raw_token: Optional[str], touch: bool = True) -> Optional[Dict[str, Any]]:
        token = str(raw_token or "").strip()
        if not token:
            return None
        token_hash = self._hash_app_session_token(token)
        with self._db_conn() as conn:
            row = conn.execute(
                """
                SELECT s.user_id, s.expires_at, u.*
                FROM auth_sessions s
                JOIN users u ON u.id = s.user_id
                WHERE s.token_hash = ?
                """,
                (token_hash,),
            ).fetchone()
            if row is None:
                return None
            expires_at = str(row["expires_at"] or "")
            if expires_at and datetime.fromisoformat(expires_at) <= _utc_now():
                conn.execute("DELETE FROM auth_sessions WHERE token_hash = ?", (token_hash,))
                conn.commit()
                return None
            if touch:
                conn.execute(
                    "UPDATE auth_sessions SET last_seen_at = ? WHERE token_hash = ?",
                    (_utc_now_iso(), token_hash),
                )
                conn.commit()
        return self._serialize_user_row(row)

    def logout_app_session(self, raw_token: Optional[str]):
        token = str(raw_token or "").strip()
        if not token:
            return
        with self._db_conn() as conn:
            conn.execute("DELETE FROM auth_sessions WHERE token_hash = ?", (self._hash_app_session_token(token),))
            conn.commit()

    def resolve_owner_context(
        self,
        session_id: Optional[str] = None,
        user_id: Optional[str] = None,
    ) -> Dict[str, Any]:
        if user_id:
            safe_user_id = _safe_owner_id(str(user_id))
            return {
                "owner_type": "user",
                "owner_key": f"user:{safe_user_id}",
                "user_id": safe_user_id,
                "session_id": _safe_session_id(session_id or "user_session_1"),
            }
        safe_session = _safe_session_id(session_id or "user_session_1")
        return {
            "owner_type": "session",
            "owner_key": f"session:{safe_session}",
            "session_id": safe_session,
            "user_id": None,
        }

    def _owner_lock(self, owner_ctx: Dict[str, Any]) -> Lock:
        owner_key = str(owner_ctx.get("owner_key") or "session:default")
        with self._owner_locks_guard:
            lock = self._owner_locks.get(owner_key)
            if lock is None:
                lock = Lock()
                self._owner_locks[owner_key] = lock
            return lock

    def _session_state_path(self, session_id: str) -> Path:
        sid = _safe_session_id(session_id)
        return MAIL_SESSION_DIR / sid / "state.json"

    def _default_state(self) -> Dict[str, Any]:
        return {
            "oauth": {},
            "whitelist": [],
            "processed_message_ids": [],
            "candidates": [],
            "last_poll_at": None,
            "updated_at": _utc_now_iso(),
        }

    def _load_state(self, session_id: str) -> Dict[str, Any]:
        state = _read_json(self._session_state_path(session_id), self._default_state())
        if not isinstance(state, dict):
            state = self._default_state()
        state.setdefault("oauth", {})
        state.setdefault("whitelist", [])
        state.setdefault("processed_message_ids", [])
        state.setdefault("candidates", [])
        state.setdefault("last_poll_at", None)
        return state

    def _save_state(self, session_id: str, state: Dict[str, Any]):
        payload = dict(state)
        payload["updated_at"] = _utc_now_iso()
        _write_json(self._session_state_path(session_id), payload)

    def _oauth_states(self) -> Dict[str, Any]:
        data = _read_json(MAIL_OAUTH_STATE_FILE, {})
        return data if isinstance(data, dict) else {}

    def _save_oauth_states(self, payload: Dict[str, Any]):
        _write_json(MAIL_OAUTH_STATE_FILE, payload)

    def _cleanup_oauth_states(self, payload: Optional[Dict[str, Any]] = None):
        states = payload if isinstance(payload, dict) else self._oauth_states()
        now = _utc_now()
        cleaned: Dict[str, Any] = {}
        for token, item in states.items():
            if not isinstance(item, dict):
                continue
            expires_at = str(item.get("expires_at") or "")
            try:
                expires_ts = datetime.fromisoformat(expires_at) if expires_at else now + timedelta(seconds=1)
            except Exception:
                expires_ts = now - timedelta(seconds=1)
            consumed_at = str(item.get("consumed_at") or "")
            if consumed_at:
                try:
                    consumed_ts = datetime.fromisoformat(consumed_at)
                except Exception:
                    consumed_ts = now - timedelta(days=1)
                if consumed_ts < now - timedelta(hours=1):
                    continue
            if expires_ts < now - timedelta(hours=1):
                continue
            cleaned[str(token)] = item
        self._save_oauth_states(cleaned)

    def _issue_oauth_state(
        self,
        flow: str,
        redirect_uri: str,
        session_id: Optional[str] = None,
        user_id: Optional[str] = None,
        scope: Optional[str] = None,
    ) -> str:
        state_token = secrets.token_urlsafe(24)
        nonce = secrets.token_urlsafe(16)
        state_map = self._oauth_states()
        now = _utc_now()
        state_map[state_token] = {
            "flow": flow,
            "session_id": _safe_session_id(session_id or "user_session_1"),
            "user_id": _safe_owner_id(user_id) if user_id else None,
            "redirect_uri": redirect_uri,
            "scope": scope or "",
            "nonce": nonce,
            "created_at": now.isoformat(),
            "expires_at": (now + timedelta(seconds=self.oauth_state_ttl_seconds)).isoformat(),
            "consumed_at": None,
        }
        self._cleanup_oauth_states(state_map)
        return state_token

    def _peek_oauth_state(self, state: str) -> Dict[str, Any]:
        payload = self._oauth_states().get(state)
        if not isinstance(payload, dict):
            raise ValueError("Invalid OAuth state.")
        return payload

    def _consume_oauth_state(self, state: str, expected_flow: str) -> Dict[str, Any]:
        state_map = self._oauth_states()
        payload = state_map.get(state)
        if not isinstance(payload, dict):
            raise ValueError("Invalid OAuth state.")
        if str(payload.get("flow") or "") != expected_flow:
            raise ValueError("OAuth flow mismatch.")
        if payload.get("consumed_at"):
            raise ValueError("OAuth state already used.")
        expires_at = str(payload.get("expires_at") or "")
        try:
            expires_ts = datetime.fromisoformat(expires_at)
        except Exception:
            raise ValueError("OAuth state expired.")
        if expires_ts <= _utc_now():
            raise ValueError("OAuth state expired.")
        payload["consumed_at"] = _utc_now_iso()
        state_map[state] = payload
        self._cleanup_oauth_states(state_map)
        return payload

    def _session_resource_dirs(self, session_id: str) -> Tuple[Path, Path, Path]:
        sid = _safe_session_id(session_id)
        root = RESOURCE_SESSION_DIR / sid
        pdf_dir = root / "pdfs"
        html_dir = root / "html"
        config_path = root / "config.json"
        pdf_dir.mkdir(parents=True, exist_ok=True)
        html_dir.mkdir(parents=True, exist_ok=True)
        if not config_path.exists():
            _write_json(config_path, {"urls": []})
        return pdf_dir, html_dir, config_path

    def _user_resource_dirs(self, user_id: str) -> Tuple[Path, Path, Path]:
        safe_user = _safe_owner_id(user_id)
        root = RESOURCE_USER_DIR / safe_user
        pdf_dir = root / "pdfs"
        html_dir = root / "html"
        config_path = root / "config.json"
        pdf_dir.mkdir(parents=True, exist_ok=True)
        html_dir.mkdir(parents=True, exist_ok=True)
        if not config_path.exists():
            _write_json(config_path, {"urls": []})
        return pdf_dir, html_dir, config_path

    def _session_add_url(self, session_id: str, url: str):
        _, _, config_path = self._session_resource_dirs(session_id)
        config = _read_json(config_path, {"urls": []})
        if not isinstance(config, dict):
            config = {"urls": []}
        config.setdefault("urls", [])
        if not any((entry.get("url") == url) for entry in config["urls"] if isinstance(entry, dict)):
            config["urls"].append({"url": url, "added_at": _utc_now_iso()})
        _write_json(config_path, config)

    def _user_add_url(self, user_id: str, url: str):
        _, _, config_path = self._user_resource_dirs(user_id)
        config = _read_json(config_path, {"urls": []})
        if not isinstance(config, dict):
            config = {"urls": []}
        config.setdefault("urls", [])
        if not any((entry.get("url") == url) for entry in config["urls"] if isinstance(entry, dict)):
            config["urls"].append({"url": url, "added_at": _utc_now_iso()})
        _write_json(config_path, config)

    def _owner_resource_dirs(self, owner_ctx: Dict[str, Any]) -> Tuple[Path, Path, Path]:
        if str(owner_ctx.get("owner_type")) == "user":
            return self._user_resource_dirs(str(owner_ctx.get("user_id") or "default"))
        return self._session_resource_dirs(str(owner_ctx.get("session_id") or "user_session_1"))

    def _owner_add_url(self, owner_ctx: Dict[str, Any], url: str):
        if str(owner_ctx.get("owner_type")) == "user":
            self._user_add_url(str(owner_ctx.get("user_id") or "default"), url)
            return
        self._session_add_url(str(owner_ctx.get("session_id") or "user_session_1"), url)

    def _default_owner_state(self) -> Dict[str, Any]:
        return self._default_state()

    def _load_user_state(self, user_id: str) -> Dict[str, Any]:
        with self._db_conn() as conn:
            row = conn.execute(
                """
                SELECT whitelist_json, processed_message_ids_json, candidates_json, last_poll_at
                FROM mail_user_state
                WHERE user_id = ?
                """,
                (str(user_id),),
            ).fetchone()
        if row is None:
            return self._default_owner_state()
        state = {
            "oauth": {},
            "whitelist": json.loads(str(row["whitelist_json"] or "[]")),
            "processed_message_ids": json.loads(str(row["processed_message_ids_json"] or "[]")),
            "candidates": json.loads(str(row["candidates_json"] or "[]")),
            "last_poll_at": row["last_poll_at"],
        }
        return state

    def _save_user_state(self, user_id: str, state: Dict[str, Any]):
        with self._db_conn() as conn:
            conn.execute(
                """
                INSERT OR REPLACE INTO mail_user_state (
                    user_id, whitelist_json, processed_message_ids_json, candidates_json, last_poll_at, updated_at
                ) VALUES (?, ?, ?, ?, ?, ?)
                """,
                (
                    str(user_id),
                    json.dumps(state.get("whitelist") or [], ensure_ascii=False),
                    json.dumps(state.get("processed_message_ids") or [], ensure_ascii=False),
                    json.dumps(state.get("candidates") or [], ensure_ascii=False),
                    state.get("last_poll_at"),
                    _utc_now_iso(),
                ),
            )
            conn.commit()

    def _load_owner_state(self, owner_ctx: Dict[str, Any]) -> Dict[str, Any]:
        if str(owner_ctx.get("owner_type")) == "user":
            return self._load_user_state(str(owner_ctx.get("user_id") or "default"))
        return self._load_state(str(owner_ctx.get("session_id") or "user_session_1"))

    def _save_owner_state(self, owner_ctx: Dict[str, Any], state: Dict[str, Any]):
        if str(owner_ctx.get("owner_type")) == "user":
            self._save_user_state(str(owner_ctx.get("user_id") or "default"), state)
            return
        self._save_state(str(owner_ctx.get("session_id") or "user_session_1"), state)

    def _get_mail_connection(self, user_id: str) -> Optional[Dict[str, Any]]:
        with self._db_conn() as conn:
            row = conn.execute("SELECT * FROM mail_connections WHERE user_id = ?", (str(user_id),)).fetchone()
        if row is None:
            return None
        return {
            "user_id": str(row["user_id"]),
            "gmail_email": str(row["gmail_email"]),
            "refresh_token": self._decrypt_secret(str(row["refresh_token_enc"] or "")),
            "access_token": str(row["access_token"] or ""),
            "access_expiry": int(row["access_expiry"] or 0),
            "scope": str(row["scope"] or self.gmail_scope),
            "connected_at": str(row["connected_at"] or ""),
            "updated_at": str(row["updated_at"] or ""),
        }

    def _save_mail_connection(
        self,
        user_id: str,
        gmail_email: str,
        refresh_token: str,
        access_token: str,
        access_expiry: int,
        scope: str,
        connected_at: Optional[str] = None,
    ):
        with self._db_conn() as conn:
            existing = conn.execute(
                "SELECT connected_at FROM mail_connections WHERE user_id = ?",
                (str(user_id),),
            ).fetchone()
            conn.execute(
                """
                INSERT OR REPLACE INTO mail_connections (
                    user_id, gmail_email, refresh_token_enc, access_token, access_expiry, scope, connected_at, updated_at
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    str(user_id),
                    str(gmail_email or "").strip().lower(),
                    self._encrypt_secret(refresh_token),
                    access_token or "",
                    int(access_expiry or 0),
                    scope or self.gmail_scope,
                    connected_at or str(existing["connected_at"] or "") or _utc_now_iso(),
                    _utc_now_iso(),
                ),
            )
            conn.commit()

    def _delete_mail_connection(self, user_id: str):
        with self._db_conn() as conn:
            conn.execute("DELETE FROM mail_connections WHERE user_id = ?", (str(user_id),))
            conn.commit()

    def _has_gmail_oauth_config(self) -> bool:
        return bool(self.gmail_client_id and self.gmail_client_secret)

    def _require_oauth_config(self):
        if not self._has_gmail_oauth_config():
            raise ValueError("Missing GOOGLE_OAUTH_CLIENT_ID/GOOGLE_OAUTH_CLIENT_SECRET.")

    def _oauth_redirect(self, redirect_uri: Optional[str]) -> str:
        uri = (redirect_uri or "").strip() or self.gmail_redirect_default
        if not uri:
            raise ValueError("Missing redirect_uri for OAuth callback.")
        return uri
    def _oauth_redirect_for_flow(self, flow: str, redirect_uri: Optional[str]) -> str:
        if flow == "app_auth":
            uri = (redirect_uri or "").strip() or self.app_auth_redirect_default
            if not uri:
                raise ValueError("Missing redirect_uri for app auth callback.")
            return uri
        return self._oauth_redirect(redirect_uri)

    def _build_google_auth_url(self, redirect_uri: str, scope: str, state_token: str) -> str:
        query = urlencode(
            {
                "client_id": self.gmail_client_id,
                "redirect_uri": redirect_uri,
                "response_type": "code",
                "scope": scope,
                "access_type": "offline" if "gmail" in scope else "online",
                "prompt": "consent" if "gmail" in scope else "select_account",
                "include_granted_scopes": "true",
                "state": state_token,
            }
        )
        return f"https://accounts.google.com/o/oauth2/v2/auth?{query}"

    def _exchange_google_oauth_code(self, code: str, redirect_uri: str) -> Dict[str, Any]:
        token_resp = requests.post(
            "https://oauth2.googleapis.com/token",
            data={
                "code": code,
                "client_id": self.gmail_client_id,
                "client_secret": self.gmail_client_secret,
                "redirect_uri": redirect_uri,
                "grant_type": "authorization_code",
            },
            timeout=30,
        )
        if token_resp.status_code >= 400:
            raise ValueError(f"OAuth token exchange failed: {token_resp.text}")
        return token_resp.json()

    def _google_get_userinfo(self, access_token: str) -> Dict[str, Any]:
        resp = requests.get(
            "https://www.googleapis.com/oauth2/v3/userinfo",
            headers={"Authorization": f"Bearer {access_token}"},
            timeout=30,
        )
        if resp.status_code >= 400:
            raise ValueError(f"Failed to read Google userinfo: {resp.text}")
        return resp.json()

    def start_app_auth(self, session_id: Optional[str], redirect_uri: Optional[str]) -> Dict[str, Any]:
        self._require_oauth_config()
        sid = _safe_session_id(session_id or "user_session_1")
        final_redirect = self._oauth_redirect_for_flow("app_auth", redirect_uri)
        state_token = self._issue_oauth_state(
            flow="app_auth",
            redirect_uri=final_redirect,
            session_id=sid,
            scope=self.app_auth_scope,
        )
        return {
            "auth_url": self._build_google_auth_url(final_redirect, self.app_auth_scope, state_token),
            "state": state_token,
        }

    def complete_app_auth(
        self,
        state: str,
        code: str,
        redirect_uri: Optional[str],
    ) -> Dict[str, Any]:
        self._require_oauth_config()
        state_payload = self._consume_oauth_state(state, "app_auth")
        final_redirect = self._oauth_redirect_for_flow(
            "app_auth",
            redirect_uri or str(state_payload.get("redirect_uri") or ""),
        )
        token_data = self._exchange_google_oauth_code(code, final_redirect)
        access_token = str(token_data.get("access_token") or "").strip()
        if not access_token:
            raise ValueError("OAuth response missing access_token.")
        profile = self._google_get_userinfo(access_token)
        user = self._upsert_user(profile)
        app_session_token = self._create_app_session(str(user["id"]))
        migration = self.migrate_session_to_user(
            user_id=str(user["id"]),
            session_id=str(state_payload.get("session_id") or ""),
        )
        return {
            "authenticated": True,
            "user": user,
            "app_session_token": app_session_token,
            "migration": migration,
        }

    def get_auth_me(self, raw_token: Optional[str]) -> Dict[str, Any]:
        user = self.get_authenticated_user(raw_token)
        return {"authenticated": bool(user), "user": user}

    def _coerce_owner_ctx(
        self,
        session_id: Optional[str] = None,
        owner_ctx: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        if owner_ctx:
            return self.resolve_owner_context(
                session_id=owner_ctx.get("session_id"),
                user_id=owner_ctx.get("user_id"),
            )
        return self.resolve_owner_context(session_id=session_id)

    def get_status(self, session_id: Optional[str] = None, owner_ctx: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        ctx = self._coerce_owner_ctx(session_id=session_id, owner_ctx=owner_ctx)
        state = self._load_owner_state(ctx)
        counts: Dict[str, int] = {"pending": 0, "applied": 0, "rejected": 0, "error": 0}
        for item in state.get("candidates") or []:
            status = str(item.get("status") or "pending").lower()
            counts[status] = counts.get(status, 0) + 1
        connection = None
        email = None
        connected = False
        if str(ctx.get("owner_type")) == "user":
            connection = self._get_mail_connection(str(ctx.get("user_id") or ""))
            connected = bool(connection and connection.get("refresh_token"))
            email = connection.get("gmail_email") if connection else None
        else:
            oauth = state.get("oauth") or {}
            connected = bool(oauth.get("refresh_token"))
            email = oauth.get("email")
        return {
            "session_id": ctx.get("session_id"),
            "user_id": ctx.get("user_id"),
            "owner_type": ctx.get("owner_type"),
            "owner_key": ctx.get("owner_key"),
            "oauth_configured": self._has_gmail_oauth_config(),
            "connected": connected,
            "email": email,
            "poll_minutes": self.poll_minutes,
            "last_poll_at": state.get("last_poll_at"),
            "whitelist_count": len(state.get("whitelist") or []),
            "candidate_counts": counts,
        }

    def begin_oauth(
        self,
        session_id: Optional[str],
        redirect_uri: Optional[str],
        owner_ctx: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        self._require_oauth_config()
        ctx = self._coerce_owner_ctx(session_id=session_id, owner_ctx=owner_ctx)
        final_redirect = self._oauth_redirect_for_flow("mail_connect", redirect_uri)
        state_token = self._issue_oauth_state(
            flow="mail_connect",
            redirect_uri=final_redirect,
            session_id=str(ctx.get("session_id") or "user_session_1"),
            user_id=str(ctx.get("user_id") or "") or None,
            scope=self.gmail_scope,
        )
        return {
            "auth_url": self._build_google_auth_url(final_redirect, self.gmail_scope, state_token),
            "state": state_token,
        }

    def complete_oauth(
        self,
        session_id: Optional[str],
        state: str,
        code: str,
        redirect_uri: Optional[str],
        owner_ctx: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        self._require_oauth_config()
        state_payload = self._consume_oauth_state(state, "mail_connect")
        ctx = self.resolve_owner_context(
            session_id=str(state_payload.get("session_id") or session_id or "user_session_1"),
            user_id=str(state_payload.get("user_id") or "") or (owner_ctx or {}).get("user_id"),
        )
        final_redirect = self._oauth_redirect_for_flow(
            "mail_connect",
            redirect_uri or str(state_payload.get("redirect_uri") or ""),
        )
        token_data = self._exchange_google_oauth_code(code, final_redirect)
        access_token = str(token_data.get("access_token") or "").strip()
        refresh_token = str(token_data.get("refresh_token") or "").strip()
        if not access_token:
            raise ValueError("OAuth response missing access_token.")

        profile = self._gmail_get_profile(access_token)
        expiry_ts = int(time.time()) + int(token_data.get("expires_in") or 3600) - 30
        if str(ctx.get("owner_type")) == "user":
            user_id = str(ctx.get("user_id") or "")
            existing = self._get_mail_connection(user_id)
            self._save_mail_connection(
                user_id=user_id,
                gmail_email=str(profile.get("emailAddress") or ""),
                refresh_token=refresh_token or str((existing or {}).get("refresh_token") or ""),
                access_token=access_token,
                access_expiry=expiry_ts,
                scope=str(token_data.get("scope") or self.gmail_scope),
            )
            return self.get_status(owner_ctx=ctx)

        state_doc = self._load_state(str(ctx.get("session_id") or "user_session_1"))
        oauth = state_doc.get("oauth") or {}
        oauth.update(
            {
                "email": profile.get("emailAddress"),
                "access_token": access_token,
                "refresh_token": refresh_token or oauth.get("refresh_token"),
                "token_type": token_data.get("token_type", "Bearer"),
                "scope": token_data.get("scope", self.gmail_scope),
                "expiry_ts": expiry_ts,
                "connected_at": _utc_now_iso(),
            }
        )
        state_doc["oauth"] = oauth
        self._save_state(str(ctx.get("session_id") or "user_session_1"), state_doc)
        return self.get_status(owner_ctx=ctx)

    def complete_oauth_from_state(
        self,
        state: str,
        code: str,
        redirect_uri: Optional[str],
    ) -> Dict[str, Any]:
        state_payload = self._peek_oauth_state(state)
        session_id = str(state_payload.get("session_id") or "user_session_1")
        owner_ctx = self.resolve_owner_context(
            session_id=session_id,
            user_id=str(state_payload.get("user_id") or "") or None,
        )
        return self.complete_oauth(
            session_id=session_id,
            state=state,
            code=code,
            redirect_uri=redirect_uri or state_payload.get("redirect_uri"),
            owner_ctx=owner_ctx,
        )

    def disconnect(self, session_id: Optional[str] = None, owner_ctx: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        ctx = self._coerce_owner_ctx(session_id=session_id, owner_ctx=owner_ctx)
        if str(ctx.get("owner_type")) == "user":
            self._delete_mail_connection(str(ctx.get("user_id") or ""))
            return self.get_status(owner_ctx=ctx)
        state = self._load_state(str(ctx.get("session_id") or "user_session_1"))
        state["oauth"] = {}
        self._save_state(str(ctx.get("session_id") or "user_session_1"), state)
        return self.get_status(owner_ctx=ctx)

    def get_whitelist(self, session_id: Optional[str] = None, owner_ctx: Optional[Dict[str, Any]] = None) -> List[str]:
        ctx = self._coerce_owner_ctx(session_id=session_id, owner_ctx=owner_ctx)
        state = self._load_owner_state(ctx)
        return [str(v).strip() for v in (state.get("whitelist") or []) if str(v).strip()]

    def set_whitelist(
        self,
        session_id: Optional[str] = None,
        senders: Optional[List[str]] = None,
        owner_ctx: Optional[Dict[str, Any]] = None,
    ) -> List[str]:
        clean = []
        for sender in senders or []:
            value = str(sender or "").strip().lower()
            if value:
                clean.append(value)
        deduped = list(dict.fromkeys(clean))
        ctx = self._coerce_owner_ctx(session_id=session_id, owner_ctx=owner_ctx)
        state = self._load_owner_state(ctx)
        state["whitelist"] = deduped
        self._save_owner_state(ctx, state)
        return deduped

    def _gmail_get_profile(self, access_token: str) -> Dict[str, Any]:
        resp = requests.get(
            "https://gmail.googleapis.com/gmail/v1/users/me/profile",
            headers={"Authorization": f"Bearer {access_token}"},
            timeout=30,
        )
        if resp.status_code >= 400:
            raise ValueError(f"Failed to read Gmail profile: {resp.text}")
        return resp.json()

    def _copy_resource_file(self, source_path: Path, target_dir: Path, source_session_id: str) -> Dict[str, Any]:
        target_dir.mkdir(parents=True, exist_ok=True)
        filename = source_path.name
        target_path = target_dir / filename
        if target_path.exists():
            source_hash = hashlib.sha1(source_path.read_bytes()).hexdigest()
            target_hash = hashlib.sha1(target_path.read_bytes()).hexdigest()
            if source_hash == target_hash:
                return {"status": "skipped_identical", "path": str(target_path)}
            stem = Path(filename).stem
            suffix = Path(filename).suffix
            target_path = target_dir / f"{stem}__from_{_safe_session_id(source_session_id)}__{source_hash[:8]}{suffix}"
        target_path.write_bytes(source_path.read_bytes())
        return {"status": "copied", "path": str(target_path)}

    def migrate_session_to_user(self, user_id: str, session_id: str) -> Dict[str, Any]:
        safe_user = _safe_owner_id(user_id)
        safe_session = _safe_session_id(session_id or "user_session_1")
        with self._owner_lock({"owner_key": f"user:{safe_user}"}):
            with self._db_conn() as conn:
                existing = conn.execute(
                    """
                    SELECT status, details_json
                    FROM mail_migration_log
                    WHERE user_id = ? AND source_session_id = ?
                    """,
                    (safe_user, safe_session),
                ).fetchone()
                if existing is not None:
                    return {
                        "status": str(existing["status"]),
                        "details": json.loads(str(existing["details_json"] or "{}")),
                    }

            session_state = self._load_state(safe_session)
            user_ctx = self.resolve_owner_context(session_id=safe_session, user_id=safe_user)
            user_state = self._load_owner_state(user_ctx)
            details: Dict[str, Any] = {
                "whitelist_added": 0,
                "candidates_added": 0,
                "processed_added": 0,
                "resources_copied": 0,
                "resource_conflicts": 0,
                "urls_added": 0,
                "mail_connection_migrated": False,
                "warnings": [],
            }

            merged_whitelist = list(
                dict.fromkeys(
                    [str(v).strip().lower() for v in (user_state.get("whitelist") or []) if str(v).strip()]
                    + [str(v).strip().lower() for v in (session_state.get("whitelist") or []) if str(v).strip()]
                )
            )
            details["whitelist_added"] = max(0, len(merged_whitelist) - len(user_state.get("whitelist") or []))
            user_state["whitelist"] = merged_whitelist

            current_processed = list(dict.fromkeys(str(v) for v in (user_state.get("processed_message_ids") or [])))
            session_processed = [str(v) for v in (session_state.get("processed_message_ids") or []) if str(v).strip()]
            combined_processed = list(dict.fromkeys(current_processed + session_processed))[-2000:]
            details["processed_added"] = max(0, len(combined_processed) - len(current_processed))
            user_state["processed_message_ids"] = combined_processed

            existing_candidate_keys = {
                (str(item.get("message_id") or ""), str(item.get("artifact_hash") or ""))
                for item in (user_state.get("candidates") or [])
            }
            merged_candidates = list(user_state.get("candidates") or [])
            for candidate in (session_state.get("candidates") or []):
                key = (str(candidate.get("message_id") or ""), str(candidate.get("artifact_hash") or ""))
                if key in existing_candidate_keys:
                    continue
                merged_candidates.append(candidate)
                existing_candidate_keys.add(key)
                details["candidates_added"] += 1
            user_state["candidates"] = merged_candidates

            session_last_poll = str(session_state.get("last_poll_at") or "")
            user_last_poll = str(user_state.get("last_poll_at") or "")
            user_state["last_poll_at"] = max(user_last_poll, session_last_poll)
            self._save_owner_state(user_ctx, user_state)

            session_oauth = session_state.get("oauth") or {}
            existing_connection = self._get_mail_connection(safe_user)
            if session_oauth.get("refresh_token") and not existing_connection:
                try:
                    self._save_mail_connection(
                        user_id=safe_user,
                        gmail_email=str(session_oauth.get("email") or ""),
                        refresh_token=str(session_oauth.get("refresh_token") or ""),
                        access_token=str(session_oauth.get("access_token") or ""),
                        access_expiry=int(session_oauth.get("expiry_ts") or 0),
                        scope=str(session_oauth.get("scope") or self.gmail_scope),
                        connected_at=str(session_oauth.get("connected_at") or _utc_now_iso()),
                    )
                    details["mail_connection_migrated"] = True
                except Exception as exc:
                    details["warnings"].append(f"mail_connection:{exc}")

            try:
                session_pdf_dir, session_html_dir, session_config = self._session_resource_dirs(safe_session)
                user_pdf_dir, user_html_dir, user_config = self._user_resource_dirs(safe_user)
                for source_dir, target_dir in ((session_pdf_dir, user_pdf_dir), (session_html_dir, user_html_dir)):
                    for source_path in source_dir.iterdir():
                        if not source_path.is_file():
                            continue
                        outcome = self._copy_resource_file(source_path, target_dir, safe_session)
                        if outcome["status"] == "copied":
                            details["resources_copied"] += 1
                        elif outcome["status"] == "skipped_identical":
                            details["resource_conflicts"] += 1
                session_config_data = _read_json(session_config, {"urls": []})
                user_config_data = _read_json(user_config, {"urls": []})
                session_urls = [
                    str(item.get("url") or "").strip()
                    for item in (session_config_data.get("urls") or [])
                    if isinstance(item, dict) and str(item.get("url") or "").strip()
                ]
                user_urls = [
                    str(item.get("url") or "").strip()
                    for item in (user_config_data.get("urls") or [])
                    if isinstance(item, dict) and str(item.get("url") or "").strip()
                ]
                merged_urls = list(dict.fromkeys(user_urls + session_urls))
                details["urls_added"] = max(0, len(merged_urls) - len(user_urls))
                _write_json(
                    user_config,
                    {"urls": [{"url": url, "added_at": _utc_now_iso()} for url in merged_urls]},
                )
            except Exception as exc:
                details["warnings"].append(f"resources:{exc}")

            status = "partial" if details["warnings"] else "success"
            with self._db_conn() as conn:
                conn.execute(
                    """
                    INSERT INTO mail_migration_log (id, user_id, source_session_id, status, details_json, created_at)
                    VALUES (?, ?, ?, ?, ?, ?)
                    """,
                    (str(uuid4()), safe_user, safe_session, status, json.dumps(details, ensure_ascii=False), _utc_now_iso()),
                )
                conn.commit()
            return {"status": status, "details": details}

    def _refresh_access_token(self, refresh_token: str) -> Dict[str, Any]:
        self._require_oauth_config()
        resp = requests.post(
            "https://oauth2.googleapis.com/token",
            data={
                "client_id": self.gmail_client_id,
                "client_secret": self.gmail_client_secret,
                "refresh_token": refresh_token,
                "grant_type": "refresh_token",
            },
            timeout=30,
        )
        if resp.status_code >= 400:
            raise ValueError(f"OAuth refresh failed: {resp.text}")
        return resp.json()

    def _ensure_access_token(
        self,
        session_id: Optional[str] = None,
        owner_ctx: Optional[Dict[str, Any]] = None,
    ) -> str:
        ctx = self._coerce_owner_ctx(session_id=session_id, owner_ctx=owner_ctx)
        now_ts = int(time.time())
        if str(ctx.get("owner_type")) == "user":
            user_id = str(ctx.get("user_id") or "")
            connection = self._get_mail_connection(user_id)
            if not connection or not connection.get("refresh_token"):
                raise ValueError("User is not connected to Gmail.")
            access_token = str(connection.get("access_token") or "")
            expiry_ts = int(connection.get("access_expiry") or 0)
            if access_token and now_ts < max(0, expiry_ts - 30):
                return access_token
            token_data = self._refresh_access_token(str(connection.get("refresh_token") or ""))
            new_access_token = str(token_data.get("access_token") or "")
            new_expiry = now_ts + int(token_data.get("expires_in") or 3600) - 30
            self._save_mail_connection(
                user_id=user_id,
                gmail_email=str(connection.get("gmail_email") or ""),
                refresh_token=str(connection.get("refresh_token") or ""),
                access_token=new_access_token,
                access_expiry=new_expiry,
                scope=str(token_data.get("scope") or connection.get("scope") or self.gmail_scope),
                connected_at=str(connection.get("connected_at") or _utc_now_iso()),
            )
            return new_access_token

        state = self._load_state(str(ctx.get("session_id") or "user_session_1"))
        oauth = state.get("oauth") or {}
        refresh_token = oauth.get("refresh_token")
        if not refresh_token:
            raise ValueError("Session is not connected to Gmail.")
        expiry_ts = int(oauth.get("expiry_ts") or 0)
        access_token = oauth.get("access_token")
        if access_token and now_ts < max(0, expiry_ts - 30):
            return access_token

        token_data = self._refresh_access_token(str(refresh_token))
        oauth["access_token"] = token_data.get("access_token")
        oauth["token_type"] = token_data.get("token_type", "Bearer")
        oauth["scope"] = token_data.get("scope", oauth.get("scope", self.gmail_scope))
        oauth["expiry_ts"] = now_ts + int(token_data.get("expires_in") or 3600) - 30
        state["oauth"] = oauth
        self._save_state(str(ctx.get("session_id") or "user_session_1"), state)
        return str(oauth.get("access_token") or "")

    def _gmail_list_messages(self, access_token: str, max_results: int = 20) -> List[Dict[str, Any]]:
        params = {"q": self.query, "maxResults": max(1, min(max_results, 100))}
        resp = requests.get(
            "https://gmail.googleapis.com/gmail/v1/users/me/messages",
            headers={"Authorization": f"Bearer {access_token}"},
            params=params,
            timeout=30,
        )
        if resp.status_code >= 400:
            raise ValueError(f"Failed to list Gmail messages: {resp.text}")
        data = resp.json()
        return list(data.get("messages") or [])

    def _gmail_get_message(self, access_token: str, message_id: str) -> Dict[str, Any]:
        resp = requests.get(
            f"https://gmail.googleapis.com/gmail/v1/users/me/messages/{message_id}",
            headers={"Authorization": f"Bearer {access_token}"},
            params={"format": "full"},
            timeout=30,
        )
        if resp.status_code >= 400:
            raise ValueError(f"Failed to fetch Gmail message {message_id}: {resp.text}")
        return resp.json()

    def _gmail_get_attachment(
        self, access_token: str, message_id: str, attachment_id: str
    ) -> bytes:
        resp = requests.get(
            f"https://gmail.googleapis.com/gmail/v1/users/me/messages/{message_id}/attachments/{attachment_id}",
            headers={"Authorization": f"Bearer {access_token}"},
            timeout=30,
        )
        if resp.status_code >= 400:
            raise ValueError(f"Failed to fetch attachment {attachment_id}: {resp.text}")
        data = resp.json()
        return _decode_b64url_bytes(str(data.get("data") or ""))

    def _extract_headers(self, payload: Dict[str, Any]) -> Dict[str, str]:
        headers: Dict[str, str] = {}
        for item in (payload.get("headers") or []):
            key = str(item.get("name") or "").strip().lower()
            if key:
                headers[key] = str(item.get("value") or "").strip()
        return headers

    def _extract_message_body_text(self, payload: Dict[str, Any]) -> str:
        blocks: List[str] = []
        for part in _iter_payload_parts(payload):
            mime = str(part.get("mimeType") or "").lower()
            body = part.get("body") or {}
            data = str(body.get("data") or "")
            if mime in {"text/plain", "text/html"} and data:
                blocks.append(_decode_b64url(data))
        return "\n".join(blocks).strip()

    def _extract_attachment_artifacts(self, payload: Dict[str, Any]) -> List[Dict[str, Any]]:
        artifacts: List[Dict[str, Any]] = []
        for part in _iter_payload_parts(payload):
            filename = str(part.get("filename") or "").strip()
            if not filename:
                continue
            ext = Path(filename).suffix.lower()
            if ext not in SUPPORTED_EXTENSIONS:
                continue
            body = part.get("body") or {}
            attachment_id = str(body.get("attachmentId") or "").strip()
            if not attachment_id:
                continue
            artifacts.append(
                {
                    "type": "attachment",
                    "name": Path(filename).name,
                    "ext": ext,
                    "mime": str(part.get("mimeType") or ""),
                    "attachment_id": attachment_id,
                    "size": int(body.get("size") or 0),
                }
            )
        return artifacts

    def _extract_links(self, text: str) -> List[str]:
        if not text:
            return []
        urls = re.findall(r"https?://[^\s<>\"]+", text)
        cleaned: List[str] = []
        for raw in urls:
            url = raw.strip(").,;\"'")
            try:
                parsed = urlparse(url)
                host = (parsed.netloc or "").lower()
                if not host:
                    continue
                if self.trusted_domains and not any(
                    host == dom or host.endswith(f".{dom}") for dom in self.trusted_domains
                ):
                    continue
                cleaned.append(url)
            except Exception:
                continue
        return list(dict.fromkeys(cleaned))

    def _sender_whitelist_check(self, sender_email: str, whitelist: List[str]) -> Tuple[bool, List[str]]:
        sender = sender_email.lower()
        domain = _extract_domain(sender)
        reasons: List[str] = []
        normalized_whitelist = [str(v).strip().lower() for v in whitelist if str(v).strip()]

        sender_allowed = True
        if normalized_whitelist:
            sender_allowed = False
            for item in normalized_whitelist:
                if "@" in item and sender == item:
                    sender_allowed = True
                    reasons.append(f"sender:{item}")
                    break
                if "@" not in item and (domain == item or domain.endswith(f".{item}")):
                    sender_allowed = True
                    reasons.append(f"domain:{item}")
                    break
        return sender_allowed, reasons

    def _rule_intent_classification(
        self,
        sender_email: str,
        subject: str,
        snippet: str,
        body: str,
        whitelist: List[str],
        attachment_count: int,
        link_count: int,
    ) -> Dict[str, Any]:
        sender_allowed, sender_reasons = self._sender_whitelist_check(sender_email, whitelist)
        reasons: List[str] = list(sender_reasons)
        if not sender_allowed:
            return {
                "mode": self.intent_mode,
                "source": "rule",
                "intent": "other",
                "confidence": 0.05,
                "is_relevant": False,
                "reasons": reasons,
                "ambiguous": False,
            }

        haystack = normalize_for_match(" ".join([subject or "", snippet or "", body or ""]))
        keyword_hits: List[str] = []
        for kw in self.relevance_keywords:
            norm_kw = normalize_for_match(kw)
            if norm_kw and norm_kw in haystack:
                keyword_hits.append(kw)

        schedule_hits: List[str] = []
        for kw in self.intent_schedule_tokens:
            norm_kw = normalize_for_match(kw)
            if norm_kw and norm_kw in haystack:
                schedule_hits.append(kw)

        registration_hits: List[str] = []
        for kw in self.intent_registration_tokens:
            norm_kw = normalize_for_match(kw)
            if norm_kw and norm_kw in haystack:
                registration_hits.append(kw)

        negative_hits: List[str] = []
        for kw in self.intent_negative_tokens:
            norm_kw = normalize_for_match(kw)
            if norm_kw and norm_kw in haystack:
                negative_hits.append(kw)

        if keyword_hits:
            reasons.append(f"keywords:{', '.join(keyword_hits[:5])}")
        if schedule_hits:
            reasons.append(f"schedule_tokens:{', '.join(schedule_hits[:5])}")
        if registration_hits:
            reasons.append(f"registration_tokens:{', '.join(registration_hits[:5])}")
        if negative_hits:
            reasons.append(f"negative_tokens:{', '.join(negative_hits[:5])}")

        score = 0.0
        score += 0.25  # sender/domain gate passed
        score += min(0.35, 0.12 * len(schedule_hits))
        score += min(0.30, 0.10 * len(registration_hits))
        score += min(0.20, 0.05 * len(keyword_hits))
        if attachment_count > 0:
            score += 0.15
        if link_count > 0:
            score += 0.05
        score -= min(0.35, 0.12 * len(negative_hits))
        score = _clip01(score)

        if schedule_hits:
            intent = "schedule_update"
        elif registration_hits:
            intent = "registration_notice"
        elif keyword_hits:
            intent = "academic_notice"
        else:
            intent = "other"

        is_relevant = score >= 0.55 and intent != "other"
        if not is_relevant and negative_hits:
            intent = "other"
        ambiguous = (0.45 <= score <= 0.72) or bool(negative_hits and (schedule_hits or registration_hits))

        return {
            "mode": self.intent_mode,
            "source": "rule",
            "intent": intent,
            "confidence": round(score, 4),
            "is_relevant": bool(is_relevant),
            "reasons": reasons,
            "ambiguous": bool(ambiguous),
        }

    def _extract_json_obj_from_text(self, text: str) -> Optional[Dict[str, Any]]:
        raw = (text or "").strip()
        if not raw:
            return None
        raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE)
        raw = re.sub(r"\s*```$", "", raw)
        try:
            obj = json.loads(raw)
            return obj if isinstance(obj, dict) else None
        except Exception:
            pass
        match = re.search(r"\{.*\}", raw, flags=re.DOTALL)
        if not match:
            return None
        try:
            obj = json.loads(match.group(0))
            return obj if isinstance(obj, dict) else None
        except Exception:
            return None

    def _llm_intent_classification(
        self,
        sender_email: str,
        subject: str,
        snippet: str,
        body: str,
        attachment_count: int,
        link_count: int,
    ) -> Optional[Dict[str, Any]]:
        if not self.gemini_api_key:
            return None
        try:
            import google.generativeai as genai  # type: ignore

            genai.configure(api_key=self.gemini_api_key)
            model = genai.GenerativeModel(self.intent_llm_model)
            prompt = (
                "Classify this email for academic schedule ingestion. "
                "Return strict JSON only with keys: intent, is_relevant, confidence, reasons.\n"
                "Allowed intent values: schedule_update, registration_notice, academic_notice, other.\n"
                f"sender_email: {sender_email}\n"
                f"subject: {subject}\n"
                f"snippet: {snippet}\n"
                f"body: {body[:5000]}\n"
                f"attachment_count: {attachment_count}\n"
                f"link_count: {link_count}\n"
            )
            response = model.generate_content(prompt, generation_config={"temperature": 0.0})
            payload = self._extract_json_obj_from_text(str(getattr(response, "text", "") or ""))
            if not payload:
                return None
            intent = str(payload.get("intent") or "other").strip().lower()
            if intent not in {"schedule_update", "registration_notice", "academic_notice", "other"}:
                intent = "other"
            is_relevant = bool(payload.get("is_relevant"))
            confidence = _clip01(float(payload.get("confidence") or 0.0))
            reasons = payload.get("reasons") if isinstance(payload.get("reasons"), list) else []
            reasons = [str(v) for v in reasons if str(v).strip()]
            return {
                "mode": self.intent_mode,
                "source": "llm",
                "intent": intent,
                "confidence": round(confidence, 4),
                "is_relevant": is_relevant,
                "reasons": reasons,
            }
        except Exception as e:
            logger.warning("Mail intent LLM classification failed: %s", e)
            return None

    def _match_relevance(
        self,
        sender_email: str,
        subject: str,
        snippet: str,
        body: str,
        whitelist: List[str],
        attachment_count: int,
        link_count: int,
    ) -> Tuple[bool, List[str], Dict[str, Any]]:
        rule_result = self._rule_intent_classification(
            sender_email=sender_email,
            subject=subject,
            snippet=snippet,
            body=body,
            whitelist=whitelist,
            attachment_count=attachment_count,
            link_count=link_count,
        )

        llm_result: Optional[Dict[str, Any]] = None
        if self.intent_mode == "llm_only":
            llm_result = self._llm_intent_classification(
                sender_email=sender_email,
                subject=subject,
                snippet=snippet,
                body=body,
                attachment_count=attachment_count,
                link_count=link_count,
            )
            final = llm_result or rule_result
        elif self.intent_mode == "hybrid":
            if bool(rule_result.get("ambiguous")):
                llm_result = self._llm_intent_classification(
                    sender_email=sender_email,
                    subject=subject,
                    snippet=snippet,
                    body=body,
                    attachment_count=attachment_count,
                    link_count=link_count,
                )
            final = llm_result or rule_result
        else:
            final = rule_result

        confidence = float(final.get("confidence") or 0.0)
        if llm_result is not None:
            is_relevant = bool(final.get("is_relevant")) and confidence >= self.intent_llm_threshold
        else:
            is_relevant = bool(final.get("is_relevant"))

        reasons = [str(v) for v in (final.get("reasons") or []) if str(v).strip()]
        classification = {
            "mode": self.intent_mode,
            "source": final.get("source", "rule"),
            "intent": final.get("intent", "other"),
            "confidence": round(confidence, 4),
            "is_relevant": bool(is_relevant),
            "reasons": reasons,
            "llm_threshold": self.intent_llm_threshold,
        }
        return bool(is_relevant), reasons, classification

    def _candidate_payload(
        self,
        message: Dict[str, Any],
        sender_email: str,
        sender_display: str,
        subject: str,
        snippet: str,
        body_text: str,
        reasons: List[str],
        artifacts: List[Dict[str, Any]],
        links: List[str],
        classification: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        links_artifacts = [{"type": "link", "url": url, "name": url} for url in links]
        artifact_hash_basis = [subject, sender_email] + [a.get("name", "") for a in artifacts] + links
        artifact_hash = hashlib.sha1("|".join(artifact_hash_basis).encode("utf-8")).hexdigest()
        intent = (classification or {}).get("intent", "other")
        confidence = (classification or {}).get("confidence", 0.0)
        snippet_text = str(snippet or "").strip()
        body_preview = " ".join(str(body_text or "").split())
        if len(body_preview) > 5000:
            body_preview = body_preview[:5000].rstrip() + "..."
        return {
            "id": str(uuid4()),
            "message_id": message.get("id"),
            "thread_id": message.get("threadId"),
            "internal_date": message.get("internalDate"),
            "sender_email": sender_email,
            "sender_display": sender_display,
            "subject": subject,
            "snippet": snippet_text,
            "body_preview": body_preview,
            "reasons": reasons,
            "status": "pending",
            "artifacts": artifacts + links_artifacts,
            "artifact_hash": artifact_hash,
            "created_at": _utc_now_iso(),
            "updated_at": _utc_now_iso(),
            "errors": [],
            "warnings": [],
            "applied_resources": [],
            "intent": intent,
            "confidence": confidence,
            "classification": classification or {},
        }

    def _cleanup_retention(self, state: Dict[str, Any]):
        cutoff = _utc_now() - timedelta(days=self.retention_days)
        keep_candidates: List[Dict[str, Any]] = []
        for item in state.get("candidates") or []:
            raw = str(item.get("created_at") or "")
            try:
                ts = datetime.fromisoformat(raw)
                if ts.tzinfo is None:
                    ts = ts.replace(tzinfo=timezone.utc)
            except Exception:
                ts = _utc_now()
            if ts >= cutoff:
                keep_candidates.append(item)
        state["candidates"] = keep_candidates
        state["processed_message_ids"] = list(dict.fromkeys((state.get("processed_message_ids") or [])[-2000:]))

    def poll_owner(self, owner_ctx: Dict[str, Any], max_messages: int = 20) -> Dict[str, Any]:
        ctx = self._coerce_owner_ctx(owner_ctx=owner_ctx)
        with self._owner_lock(ctx):
            access_token = self._ensure_access_token(owner_ctx=ctx)
            state = self._load_owner_state(ctx)
            processed = set(str(v) for v in (state.get("processed_message_ids") or []))
            existing_pairs = {
                (str(c.get("message_id") or ""), str(c.get("artifact_hash") or ""))
                for c in (state.get("candidates") or [])
            }
            whitelist = state.get("whitelist") or []

            messages = self._gmail_list_messages(access_token, max_results=max_messages)
            inspected = 0
            created = 0
            for ref in messages:
                message_id = str(ref.get("id") or "").strip()
                if not message_id or message_id in processed:
                    continue
                inspected += 1
                try:
                    msg = self._gmail_get_message(access_token, message_id)
                except Exception as e:
                    logger.warning("Failed to fetch Gmail message %s: %s", message_id, e)
                    processed.add(message_id)
                    continue

                payload = msg.get("payload") or {}
                headers = self._extract_headers(payload)
                subject = headers.get("subject", "")
                sender_display = headers.get("from", "")
                sender_email = _extract_email(sender_display)
                snippet = str(msg.get("snippet") or "")
                body_text = self._extract_message_body_text(payload)
                attachments = self._extract_attachment_artifacts(payload)
                links = self._extract_links("\n".join([snippet, body_text]))
                relevant, reasons, classification = self._match_relevance(
                    sender_email=sender_email,
                    subject=subject,
                    snippet=snippet,
                    body=body_text,
                    whitelist=whitelist,
                    attachment_count=len(attachments),
                    link_count=len(links),
                )

                if relevant and (attachments or links):
                    candidate = self._candidate_payload(
                        message=msg,
                        sender_email=sender_email,
                        sender_display=sender_display,
                        subject=subject,
                        snippet=snippet,
                        body_text=body_text,
                        reasons=reasons,
                        artifacts=attachments,
                        links=links,
                        classification=classification,
                    )
                    key = (str(candidate.get("message_id") or ""), str(candidate.get("artifact_hash") or ""))
                    if key not in existing_pairs:
                        state["candidates"].append(candidate)
                        existing_pairs.add(key)
                        created += 1
                processed.add(message_id)

            state["processed_message_ids"] = list(processed)
            state["last_poll_at"] = _utc_now_iso()
            self._cleanup_retention(state)
            self._save_owner_state(ctx, state)
            return {
                "session_id": ctx.get("session_id"),
                "user_id": ctx.get("user_id"),
                "owner_type": ctx.get("owner_type"),
                "inspected_messages": inspected,
                "new_candidates": created,
                "last_poll_at": state.get("last_poll_at"),
            }

    def poll_session(self, session_id: str, max_messages: int = 20) -> Dict[str, Any]:
        return self.poll_owner(self.resolve_owner_context(session_id=session_id), max_messages=max_messages)

    def poll_all_connected_users(self) -> Dict[str, Any]:
        results: List[Dict[str, Any]] = []
        with self._db_conn() as conn:
            rows = conn.execute("SELECT user_id FROM mail_connections").fetchall()
        if not rows:
            return {"polled": 0, "results": results}
        for row in rows:
            user_id = str(row["user_id"])
            try:
                status = self.get_status(owner_ctx=self.resolve_owner_context(user_id=user_id))
                if not status.get("connected"):
                    continue
                results.append(self.poll_owner(self.resolve_owner_context(user_id=user_id)))
            except Exception as e:
                results.append({"user_id": user_id, "error": str(e)})
        return {"polled": len(results), "results": results}

    def poll_all_connected_sessions(self) -> Dict[str, Any]:
        return self.poll_all_connected_users()

    def list_candidates(
        self,
        session_id: Optional[str] = None,
        status: Optional[str] = None,
        owner_ctx: Optional[Dict[str, Any]] = None,
    ) -> List[Dict[str, Any]]:
        ctx = self._coerce_owner_ctx(session_id=session_id, owner_ctx=owner_ctx)
        state = self._load_owner_state(ctx)
        candidates = list(state.get("candidates") or [])
        if status:
            expected = status.strip().lower()
            candidates = [item for item in candidates if str(item.get("status") or "").lower() == expected]
        candidates.sort(key=lambda item: str(item.get("created_at") or ""), reverse=True)
        return candidates

    def reject_candidate(
        self,
        session_id: Optional[str] = None,
        candidate_id: str = "",
        reason: Optional[str] = None,
        owner_ctx: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        ctx = self._coerce_owner_ctx(session_id=session_id, owner_ctx=owner_ctx)
        with self._owner_lock(ctx):
            state = self._load_owner_state(ctx)
            for item in state.get("candidates") or []:
                if str(item.get("id")) != str(candidate_id):
                    continue
                item["status"] = "rejected"
                item["rejected_reason"] = reason or "manual_reject"
                item["updated_at"] = _utc_now_iso()
                self._save_owner_state(ctx, state)
                return item
        raise ValueError("Candidate not found.")

    def _docx_to_html(self, content: bytes) -> str:
        try:
            from docx import Document  # type: ignore
        except Exception as e:
            raise RuntimeError(f"DOCX parser unavailable: {e}")

        doc = Document(BytesIO(content))
        parts: List[str] = ["<html><body>"]
        for para in doc.paragraphs:
            text = (para.text or "").strip()
            if text:
                parts.append(f"<p>{text}</p>")
        for table in doc.tables:
            parts.append("<table border='1'>")
            for row in table.rows:
                parts.append("<tr>")
                for cell in row.cells:
                    parts.append(f"<td>{(cell.text or '').strip()}</td>")
                parts.append("</tr>")
            parts.append("</table>")
        parts.append("</body></html>")
        return "\n".join(parts)

    def _xlsx_to_html(self, content: bytes) -> str:
        try:
            from openpyxl import load_workbook  # type: ignore
        except Exception as e:
            raise RuntimeError(f"XLSX parser unavailable: {e}")

        wb = load_workbook(BytesIO(content), data_only=True)
        parts: List[str] = ["<html><body>"]
        for sheet in wb.worksheets:
            parts.append(f"<h3>{sheet.title}</h3>")
            parts.append("<table border='1'>")
            for row in sheet.iter_rows(values_only=True):
                parts.append("<tr>")
                for cell in row:
                    value = "" if cell is None else str(cell)
                    parts.append(f"<td>{value}</td>")
                parts.append("</tr>")
            parts.append("</table>")
        parts.append("</body></html>")
        return "\n".join(parts)

    def _save_attachment_as_resource(
        self,
        owner_ctx: Dict[str, Any],
        artifact: Dict[str, Any],
        file_content: bytes,
    ) -> Tuple[Optional[str], Optional[str], Optional[str]]:
        pdf_dir, html_dir, _ = self._owner_resource_dirs(owner_ctx)
        filename = Path(str(artifact.get("name") or "attachment")).name
        ext = Path(filename).suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            return None, None, f"Unsupported extension: {ext}"

        if ext == ".pdf":
            target = pdf_dir / filename
            target.write_bytes(file_content)
            return str(target), "pdf", None
        if ext in {".html", ".htm"}:
            target = html_dir / filename
            target.write_bytes(file_content)
            return str(target), "html", None
        if ext == ".docx":
            html_content = self._docx_to_html(file_content)
            target = html_dir / f"{Path(filename).stem}.docx.html"
            target.write_text(html_content, encoding="utf-8")
            return str(target), "html", None
        if ext in {".xlsx", ".xls"}:
            html_content = self._xlsx_to_html(file_content)
            target = html_dir / f"{Path(filename).stem}.xlsx.html"
            target.write_text(html_content, encoding="utf-8")
            return str(target), "html", None
        return None, None, f"Unsupported extension: {ext}"

    def apply_candidate(
        self,
        session_id: Optional[str] = None,
        candidate_id: str = "",
        owner_ctx: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        ctx = self._coerce_owner_ctx(session_id=session_id, owner_ctx=owner_ctx)
        with self._owner_lock(ctx):
            state = self._load_owner_state(ctx)
            target: Optional[Dict[str, Any]] = None
            for item in state.get("candidates") or []:
                if str(item.get("id")) == str(candidate_id):
                    target = item
                    break
            if target is None:
                raise ValueError("Candidate not found.")
            if str(target.get("status") or "").lower() == "applied":
                return target

            access_token = self._ensure_access_token(owner_ctx=ctx)
            message_id = str(target.get("message_id") or "")
            errors: List[str] = []
            warnings: List[str] = []
            applied: List[Dict[str, Any]] = []

            for artifact in (target.get("artifacts") or []):
                art_type = str(artifact.get("type") or "")
                try:
                    if art_type == "link":
                        url = str(artifact.get("url") or "").strip()
                        if not url:
                            continue
                        self._owner_add_url(ctx, url)
                        applied.append({"type": "url", "name": url})
                        continue

                    if art_type != "attachment":
                        continue

                    attachment_id = str(artifact.get("attachment_id") or "").strip()
                    if not attachment_id or not message_id:
                        warnings.append(f"Skip attachment {artifact.get('name')}: missing attachment_id/message_id")
                        continue
                    payload = self._gmail_get_attachment(access_token, message_id, attachment_id)
                    saved_path, saved_type, err = self._save_attachment_as_resource(ctx, artifact, payload)
                    if err:
                        warnings.append(f"{artifact.get('name')}: {err}")
                        continue
                    applied.append({"type": saved_type, "name": artifact.get("name"), "path": saved_path})
                except Exception as e:
                    errors.append(f"{artifact.get('name')}: {e}")

            target["applied_resources"] = applied
            target["errors"] = errors
            target["warnings"] = warnings
            target["status"] = "applied" if applied and not errors else ("error" if errors else "rejected")
            target["updated_at"] = _utc_now_iso()
            if target["status"] == "applied":
                target["applied_at"] = _utc_now_iso()
            self._cleanup_retention(state)
            self._save_owner_state(ctx, state)
            return target


mail_agent_service = MailAgentService()
